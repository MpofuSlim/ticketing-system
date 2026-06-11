"""Generate the merchant POS integration guide as a styled .docx.

Run:  python3 docs/integration/build_merchant_integration_doc.py
Output: docs/integration/InnBucks-Merchant-POS-Integration-Guide.docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# InnBucks brand — sampled from the official square logo (dark navy bg, four
# accent dots, white wordmark). Used consistently across this doc so the print
# version reads as InnBucks corporate, not generic Word.
BRAND_NAVY = RGBColor(0x1F, 0x24, 0x40)        # logo background
BRAND_BLUE = RGBColor(0x1F, 0x24, 0x40)        # primary heading colour
DOT_YELLOW = RGBColor(0xE9, 0xB7, 0x21)
DOT_PURPLE = RGBColor(0x8E, 0x4F, 0xB7)
DOT_TEAL   = RGBColor(0x2A, 0xA6, 0x84)
DOT_RED    = RGBColor(0xD8, 0x3A, 0x3A)
ACCENT_GOLD = DOT_YELLOW                       # accent used for HTTP verbs
TEXT_DARK = RGBColor(0x1F, 0x29, 0x37)
TEXT_MUTED = RGBColor(0x65, 0x70, 0x80)
CODE_BG = RGBColor(0xF3, 0xF4, 0xF6)
TABLE_HEADER = RGBColor(0xE6, 0xEC, 0xF5)

LOGO_PATH = Path(__file__).resolve().parent / "innbucks-logo.png"


def add_run(p, text, *, bold=False, italic=False, color=None, size=None, font="Calibri"):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def code_block(doc, lines):
    """Render a fenced code block (monospace, light grey background)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F3F4F6")
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, line, color=TEXT_DARK, size=9, font="Consolas")
    doc.add_paragraph()  # breathing room after


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return h


def body(doc, text, *, bold_first_n=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_first_n:
        bold_part = text[:bold_first_n]
        rest = text[bold_first_n:]
        add_run(p, bold_part, bold=True, color=TEXT_DARK, size=10.5)
        add_run(p, rest, color=TEXT_DARK, size=10.5)
    else:
        add_run(p, text, color=TEXT_DARK, size=10.5)
    return p


def endpoint_row(table, method, path, desc, role):
    row = table.add_row().cells
    row[0].text = method
    row[1].text = path
    row[2].text = desc
    row[3].text = role
    for c in row:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = TEXT_DARK
    # colour the method cell
    method_run = row[0].paragraphs[0].runs[0]
    method_run.bold = True
    method_run.font.color.rgb = BRAND_BLUE if method == "GET" else ACCENT_GOLD
    # monospace the path
    path_run = row[1].paragraphs[0].runs[0]
    path_run.font.name = "Consolas"
    path_run.font.size = Pt(9.5)


def endpoint_summary_table(doc, rows):
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Path"
    hdr[2].text = "Purpose"
    hdr[3].text = "Required role"
    for c in hdr:
        shade_cell(c, "E6ECF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = BRAND_BLUE
    for r in rows:
        endpoint_row(t, *r)
    doc.add_paragraph()


def kv_table(doc, rows):
    """Two-column 'field/notes' table for payload field references."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Notes"
    for c in hdr:
        shade_cell(c, "E6ECF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = BRAND_BLUE
                r.font.size = Pt(10)
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
        # mono-font the field name
        for r in row[0].paragraphs[0].runs:
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = TEXT_DARK
        for r in row[1].paragraphs[0].runs:
            r.font.size = Pt(10)
            r.font.color.rgb = TEXT_DARK
    doc.add_paragraph()


# ----------------------------------------------------------------- build doc
doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Set default Calibri body
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ----------------------- Cover page
# Logo (if present at docs/integration/innbucks-logo.png) — otherwise a
# styled text header in the brand palette as a graceful fallback.
if LOGO_PATH.exists():
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_p.add_run()
    logo_run.add_picture(str(LOGO_PATH), width=Cm(3.5))
else:
    # Four-dot brand mark in text — close to the visual identity even without the PNG.
    dots_p = doc.add_paragraph()
    dots_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(dots_p, "● ", color=DOT_YELLOW, size=22, bold=True)
    add_run(dots_p, "● ", color=DOT_PURPLE, size=22, bold=True)
    add_run(dots_p, "● ", color=DOT_TEAL,   size=22, bold=True)
    add_run(dots_p, "●",  color=DOT_RED,    size=22, bold=True)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(cover, "InnBucks", bold=True, color=BRAND_NAVY, size=42)
add_run(cover, "®", color=BRAND_NAVY, size=14)

strip = doc.add_paragraph()
add_run(strip, "Loyalty Platform", color=TEXT_MUTED, size=14, italic=True)

# Spacing
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(title, "Merchant POS Integration Guide", bold=True, color=BRAND_BLUE, size=24)

sub = doc.add_paragraph()
add_run(sub, "How to enrol customers, post transactions, accrue points, issue and redeem vouchers — end to end.",
        color=TEXT_MUTED, size=11, italic=True)

# Document meta block
for _ in range(2):
    doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.autofit = True
labels = [
    ("Version", "1.0"),
    ("Date", date.today().isoformat()),
    ("Audience", "Merchant POS / shop-system integrators"),
    ("Classification", "Confidential — for partner use only"),
]
for i, (k, v) in enumerate(labels):
    meta.rows[i].cells[0].text = ""
    meta.rows[i].cells[1].text = ""
    p0 = meta.rows[i].cells[0].paragraphs[0]
    add_run(p0, k, bold=True, color=BRAND_BLUE, size=10)
    p1 = meta.rows[i].cells[1].paragraphs[0]
    add_run(p1, v, color=TEXT_DARK, size=10)

for _ in range(3):
    doc.add_paragraph()

prep = doc.add_paragraph()
add_run(prep, "Prepared by", bold=True, color=BRAND_BLUE, size=10)
prep.add_run().add_break()
add_run(prep, "Tawanda Mpofu", bold=True, color=TEXT_DARK, size=14)
prep.add_run().add_break()
add_run(prep, "Senior Backend Engineer  ·  Innovation Team", color=TEXT_MUTED, size=10.5)
prep.add_run().add_break()
add_run(prep, "tmpofu@innbucks.co.zw", color=BRAND_BLUE, size=10.5)

doc.add_page_break()

# ----------------------- 1. Overview
heading(doc, "1. Overview", 1)
body(doc,
     "InnBucks Loyalty is a multi-tenant points-and-vouchers platform. Each tenant "
     "(your brand) owns one or more merchants; each merchant owns one or more shops "
     "(physical outlets, kiosks, or e-commerce storefronts). Every customer transaction "
     "your POS reports earns the customer points according to your published rules, "
     "and customers can redeem points or vouchers back at the till in the same call.")

body(doc,
     "This document walks through the integration end to end from a merchant POS / shop-system "
     "perspective: the one-time setup, the per-transaction call you'll make on every sale, "
     "and the voucher / wallet lookups you'll need to surface to the customer at checkout.",
     bold_first_n=0)

body(doc, "There are three things to integrate against:", bold_first_n=0)
ov = doc.add_paragraph(style="List Bullet")
add_run(ov, "Customer identity — ", bold=True, color=TEXT_DARK, size=10.5)
add_run(ov, "the customer's phone number is the universal key. You enrol them once "
            "at tier 1, then every subsequent call references that phone.", color=TEXT_DARK, size=10.5)
ov2 = doc.add_paragraph(style="List Bullet")
add_run(ov2, "Transactions (earn + redeem) — ", bold=True, color=TEXT_DARK, size=10.5)
add_run(ov2, "a single POST per sale handles cash, points, or a mix. Points earned "
             "(rules + campaign multipliers) and points burned commit atomically.",
        color=TEXT_DARK, size=10.5)
ov3 = doc.add_paragraph(style="List Bullet")
add_run(ov3, "Vouchers — ", bold=True, color=TEXT_DARK, size=10.5)
add_run(ov3, "issue, list, and redeem. Each voucher carries an HMAC-signed code that "
             "the platform verifies on every present-at-till.", color=TEXT_DARK, size=10.5)

doc.add_paragraph()

# Endpoint summary
heading(doc, "1.1 Endpoint reference at a glance", 2)
body(doc, "All paths in this guide are relative to the gateway base URL. Send every authenticated "
          "request with the Authorization header you receive from /auth/login.")
endpoint_summary_table(doc, [
    ("POST", "/auth/login", "Obtain a JWT for the merchant operator", "Anonymous"),
    ("POST", "/auth/customer/register", "Enrol a new customer (tier 1)", "Anonymous"),
    ("POST", "/payments/shop-checkout", "Per-sale call — earn / redeem points", "SHOP_ADMIN"),
    ("GET",  "/loyalty/vouchers/users/by-phone/{phone}/active", "List a customer's active vouchers", "SHOP_ADMIN"),
    ("POST", "/loyalty/vouchers/issue", "Manually issue a voucher to a customer", "MERCHANT_ADMIN"),
    ("POST", "/loyalty/vouchers/redeem", "Redeem a voucher at the till", "SHOP_ADMIN"),
    ("GET",  "/loyalty/users/me/wallet", "Customer-facing wallet balance", "CUSTOMER"),
    ("POST", "/loyalty/qr/issue", "Mint a QR code for a voucher (mobile)", "MERCHANT_ADMIN"),
    ("POST", "/loyalty/qr/consume", "Consume a scanned voucher QR code", "SHOP_ADMIN"),
])

# ----------------------- 2. Authentication
doc.add_page_break()
heading(doc, "2. Authentication", 1)
body(doc, "Every call goes through the API gateway over HTTPS. The gateway terminates TLS, "
          "validates the JWT issued at login, and forwards the call to the matching service. "
          "There are three identity types you'll work with:")

body(doc, "MERCHANT_ADMIN — your brand's head office account. Used for one-time setup "
          "(creating merchants, shops, rules, voucher templates) and back-office tasks.",
     bold_first_n=15)
body(doc, "SHOP_ADMIN — the till operator's account. Used for every per-sale call: "
          "shop-checkout, voucher redeem, QR consume. The SHOP_ADMIN's JWT is scoped to "
          "one shop, so the customer cannot accidentally have their points credited to "
          "another outlet.",
     bold_first_n=10)
body(doc, "CUSTOMER — issued to the buyer themselves when they log in to the InnBucks "
          "customer app. You only need this if you also operate the customer-facing app; "
          "the POS itself never uses a CUSTOMER token.",
     bold_first_n=8)

heading(doc, "2.1 Obtaining a token", 2)
body(doc, "Submit the operator's email + password to /auth/login:")
code_block(doc, [
    "POST /auth/login",
    "Content-Type: application/json",
    "",
    "{",
    '  "identifier": "till-1@chickeninn.co.zw",',
    '  "password":   "•••••••••"',
    "}",
])
body(doc, "Successful response (200 OK):")
code_block(doc, [
    "{",
    '  "code": "200 OK",',
    '  "message": "Login successful",',
    '  "data": {',
    '    "token":        "eyJhbGciOiJIUzI1NiJ9…",',
    '    "refreshToken": "eyJhbGciOiJIUzI1NiJ9…",',
    '    "email":        "till-1@chickeninn.co.zw",',
    '    "roles":        ["SHOP_ADMIN"],',
    '    "defaultServices": ["loyalty"],',
    '    "mfaRequired":  false,',
    '    "mustChangePassword": false',
    "  }",
    "}",
])
body(doc, "Use the access token as Authorization: Bearer <token> on every subsequent call. "
          "Tokens expire after 24 hours by default; call /auth/refresh with the refreshToken "
          "to mint a new pair without prompting the user to log in again.")

# ----------------------- 3. Hierarchy
heading(doc, "3. The platform hierarchy", 1)
body(doc, "Everything you create at setup time hangs off a tenant. The hierarchy is fixed:")
code_block(doc, [
    "Tenant            ← your brand (e.g. \"Chicken Inn\")",
    "  └─ Merchant     ← a business unit (e.g. \"Chicken Inn Restaurants\")",
    "       └─ Shop    ← a physical outlet (e.g. \"Avondale Branch\")",
    "             └─ Transactions belong here; customers earn/burn against this shopId",
])
body(doc, "Voucher templates and earn rules are set at the merchant level — they apply across "
          "every shop owned by that merchant. Campaign multipliers can be scoped narrower (e.g. "
          "\"double points on Friday at Avondale only\").")

# ----------------------- 4. One-time setup
doc.add_page_break()
heading(doc, "4. One-time setup (head office, MERCHANT_ADMIN)", 1)
body(doc, "Done once when you onboard. After this, your POS only needs the per-sale call in §6.")

heading(doc, "4.1 Create a merchant", 2)
code_block(doc, [
    "POST /loyalty/merchants",
    "Authorization: Bearer <MERCHANT_ADMIN token>",
    "",
    "{",
    '  "name":           "Chicken Inn Restaurants",',
    '  "registrationNo": "RC-2017-04123",',
    '  "contactEmail":   "ops@chickeninn.co.zw",',
    '  "contactPhone":   "+263242123456",',
    '  "address":        "12 Speke Avenue, Harare"',
    "}",
])
body(doc, "Response carries the merchantId — store it; you'll need it for shops, rules, and templates.")

heading(doc, "4.2 Create each shop", 2)
code_block(doc, [
    "POST /loyalty/shops",
    "",
    "{",
    '  "merchantId":   "b4c0d2e3-2345-6789-abcd-ef0123456789",',
    '  "name":         "Avondale Branch",',
    '  "address":      "King George Rd, Avondale, Harare",',
    '  "phoneNumber":  "+263242456789"',
    "}",
])
body(doc, "Each shop's response carries a shopId — that's what your POS will send on every "
          "shop-checkout call. If you have many outlets, /loyalty/shops/bulk-upload accepts a "
          "CSV; ask InnBucks for the template.")

heading(doc, "4.3 Define an earn rule", 2)
body(doc, "Rules say \"for every $X of cash spent, the customer earns Y points.\" The default "
          "rule below earns 1.25 points per US dollar at any of the merchant's shops:")
code_block(doc, [
    "POST /loyalty/rules",
    "",
    "{",
    '  "merchantId":     "b4c0d2e3-2345-6789-abcd-ef0123456789",',
    '  "name":           "Standard earn",',
    '  "pointsPerUnit":  1.25,',
    '  "minSpend":       0.00,',
    '  "active":         true',
    "}",
])
body(doc, "You can stack a campaign on top of the base rule (e.g. \"2× points on Mondays\") "
          "via POST /loyalty/rules/campaigns. The platform automatically applies the active "
          "campaign multiplier at shop-checkout time.")

heading(doc, "4.4 Create voucher templates (optional but recommended)", 2)
body(doc, "A voucher template defines the kind of voucher customers can earn or that you can "
          "issue. Each template names the discount amount, validity period, and shop scope.")
code_block(doc, [
    "POST /loyalty/vouchers/templates",
    "",
    "{",
    '  "merchantId":     "b4c0d2e3-2345-6789-abcd-ef0123456789",',
    '  "name":           "$5 OFF main meal",',
    '  "type":           "SINGLE_USE",',
    '  "discountType":   "FIXED",',
    '  "discountValue":  5.00,',
    '  "validityDays":   30,',
    '  "appliesToShops": []          // empty = every shop owned by the merchant',
    "}",
])
body(doc, "Type can be SINGLE_USE (one redemption), MULTI_USE (up to N uses), CAMPAIGN "
          "(time-windowed promo), or REFERRAL (issued when a customer refers another).")

# ----------------------- 5. Customer enrolment
doc.add_page_break()
heading(doc, "5. Customer enrolment at the till", 1)
body(doc, "The customer's phone number is their identity across InnBucks. If they're new "
          "to the platform, register them once at tier 1 — this takes <5 seconds at the till.")
code_block(doc, [
    "POST /auth/customer/register",
    "Content-Type: application/json",
    "",
    "{",
    '  "firstName":   "Tanaka",',
    '  "lastName":    "Moyo",',
    '  "phoneNumber": "+263771234567",',
    '  "country":     "Zimbabwe"',
    "}",
])
body(doc, "Successful response (201 Created):")
code_block(doc, [
    "{",
    '  "code": "201 CREATED",',
    '  "message": "Registration successful.",',
    '  "data": {',
    '    "token":     "eyJhbGciOiJIUzI1NiJ9…",   // customer token, not the POS token',
    '    "roles":     ["CUSTOMER"],',
    '    "tier":      1,',
    '    "verified":  false,',
    '    "nextStep":  "Verify phone at /auth/otp/verify, then submit personal details at',
    '                 /auth/customer/register/tier2"',
    "  }",
    "}",
])
body(doc, "The till does not have to do anything with the returned token — the customer's "
          "wallet is created the moment they're enrolled, and your subsequent shop-checkout "
          "calls reference them by phone number. The token is for when the customer downloads "
          "the InnBucks customer app.")
body(doc, "If the phone is already registered, you'll receive HTTP 400 with "
          "\"Phone already registered\". That's fine — skip straight to step 6.")

# ----------------------- 6. The transaction call
heading(doc, "6. Per-sale: POST /payments/shop-checkout", 1)
body(doc, "This is the single most important endpoint in the integration. It handles both "
          "earn (cash leg) and burn (points leg) atomically. Call it once per sale, regardless "
          "of how the customer paid.")

heading(doc, "6.1 Request", 2)
code_block(doc, [
    "POST /payments/shop-checkout",
    "Authorization: Bearer <SHOP_ADMIN token>",
    "Content-Type: application/json",
    "X-Customer-Phone: +263771234567        // header: identifies the buyer",
    "",
    "{",
    '  "shopId":         "5b1c2d3e-4567-890a-bcde-f01234567890",',
    '  "paymentMethod":  "CASH_AND_POINTS",',
    '  "cashAmount":     10.00,    // USD',
    '  "pointsAmount":   200.00    // points to burn',
    "}",
])
kv_table(doc, [
    ("shopId", "Required. The outlet where the sale is taking place. Must belong to a merchant in your tenant."),
    ("paymentMethod", "Required. CASH (cash-only), POINTS (points-only), or CASH_AND_POINTS (mixed)."),
    ("cashAmount", "Cash portion in USD. Must be > 0 for CASH and CASH_AND_POINTS; omitted or 0 for POINTS."),
    ("pointsAmount", "Points to spend from the customer's wallet. Must be > 0 for POINTS and CASH_AND_POINTS; omitted or 0 for CASH."),
])
body(doc, "Notes on the request:")
n1 = doc.add_paragraph(style="List Bullet")
add_run(n1, "The cash amount is informational. ", bold=True, color=TEXT_DARK, size=10.5)
add_run(n1, "InnBucks does not collect the cash — the merchant settles that at the till. "
            "The platform uses the cash figure only to compute points earned.",
        color=TEXT_DARK, size=10.5)
n2 = doc.add_paragraph(style="List Bullet")
add_run(n2, "The customer's phone number is sent in the X-Customer-Phone header, ", bold=True, color=TEXT_DARK, size=10.5)
add_run(n2, "not in the body — keeps it consistent with our other endpoints and lets the "
            "platform audit it independently of the request body.",
        color=TEXT_DARK, size=10.5)
n3 = doc.add_paragraph(style="List Bullet")
add_run(n3, "No reference field. ", bold=True, color=TEXT_DARK, size=10.5)
add_run(n3, "The platform generates the idempotency reference automatically (returned on the "
            "response as SHOP-<uuid>). Print it on the customer's receipt.",
        color=TEXT_DARK, size=10.5)

heading(doc, "6.2 Response (200 OK)", 2)
code_block(doc, [
    "{",
    '  "code": "200 OK",',
    '  "message": "Shop checkout processed successfully",',
    '  "data": {',
    '    "transactionId":      "f0e1d2c3-4567-890a-bcde-f01234567890",',
    '    "shopId":             "5b1c2d3e-4567-890a-bcde-f01234567890",',
    '    "merchantId":         "b4c0d2e3-2345-6789-abcd-ef0123456789",',
    '    "msisdn":             "+263771234567",',
    '    "paymentMethod":      "CASH_AND_POINTS",',
    '    "cashAmount":         10.00,',
    '    "pointsRedeemed":     200.0000,',
    '    "pointsEarned":       12.5000,    // rules × campaign multiplier',
    '    "walletBalanceAfter": 1812.5000,',
    '    "processedAt":        "2026-06-11T10:30:00",',
    '    "reference":          "SHOP-7c9e6679-7425-40de-944b-e07fc1f90ae7"',
    "  }",
    "}",
])
body(doc, "Print pointsEarned and walletBalanceAfter on the receipt — that's the loyalty value "
          "the customer just got. The reference is your idempotency key; if you ever retry the "
          "same physical sale, send a different X-Idempotency-Key (it's the platform that "
          "generates the reference on its end, but your POS must dedupe its own retries).")

heading(doc, "6.3 Failure responses", 2)
body(doc, "Three categories of failure your POS must handle:")
code_block(doc, [
    "# Validation failure (400 BAD_REQUEST)",
    "{",
    '  "code": "400 BAD_REQUEST",',
    '  "message": "paymentMethod=CASH requires cashAmount > 0 and pointsAmount must be null/zero",',
    '  "data": null',
    "}",
    "",
    "# Insufficient points balance (400 BAD_REQUEST, from loyalty-service)",
    "{",
    '  "code": "400 BAD_REQUEST",',
    '  "message": "Insufficient wallet balance: requested 200 had 145",',
    '  "data": null',
    "}",
    "",
    "# Inactive shop or merchant (400 BAD_REQUEST)",
    "{",
    '  "code": "400 BAD_REQUEST",',
    '  "message": "Merchant is not active",',
    '  "data": null',
    "}",
    "",
    "# Loyalty service temporarily unreachable (503 SERVICE_UNAVAILABLE)",
    "{",
    '  "code": "503 SERVICE_UNAVAILABLE",',
    '  "message": "Unable to reach loyalty-service for checkout",',
    '  "data": null',
    "}",
])
body(doc, "On 503, retry once after ~500ms. On any 4xx, surface the message to the cashier; "
          "do NOT retry — the call already returned its terminal answer.")

# ----------------------- 7. Vouchers
doc.add_page_break()
heading(doc, "7. Voucher lifecycle", 1)
body(doc, "Vouchers are how customers get one-off discounts (e.g. \"$5 OFF main meal\"). They "
          "can be issued by the merchant directly (CRM-style), earned automatically when a "
          "customer hits a points threshold, or distributed via a campaign code. At the till, "
          "your POS does two things with vouchers: list them and redeem them.")

heading(doc, "7.1 List a customer's active vouchers", 2)
body(doc, "Before the customer's order is finalised, you can show them the discounts they "
          "already hold so they don't forget:")
code_block(doc, [
    "GET /loyalty/vouchers/users/by-phone/+263771234567/active",
    "Authorization: Bearer <SHOP_ADMIN token>",
])
code_block(doc, [
    "{",
    '  "code": "200 OK",',
    '  "message": "Active vouchers retrieved",',
    '  "data": [',
    '    {',
    '      "voucherId":    "c1b7e9f0-9012-3456-0123-456789012345",',
    '      "code":         "VCH-AB12-CD34-EF56",',
    '      "templateName": "$5 OFF main meal",',
    '      "discountType": "FIXED",',
    '      "discountValue": 5.00,',
    '      "status":       "ISSUED",',
    '      "usesRemaining": 1,',
    '      "expiresAt":    "2026-07-08T23:59:59"',
    '    },',
    '    {',
    '      "voucherId":    "d2a8f0e1-0123-4567-1234-567890123456",',
    '      "code":         "VCH-XY99-ZZ44-MM11",',
    '      "templateName": "10% OFF chicken bucket",',
    '      "discountType": "PERCENT",',
    '      "discountValue": 10.00,',
    '      "status":       "VIEWED",',
    '      "usesRemaining": 1,',
    '      "expiresAt":    "2026-06-30T23:59:59"',
    '    }',
    "  ]",
    "}",
])

heading(doc, "7.2 Redeem a voucher at the till", 2)
body(doc, "When the customer chooses a voucher to apply (either by stating its code or by "
          "scanning their app's QR code), submit a redeem call before posting the cash sale "
          "via shop-checkout:")
code_block(doc, [
    "POST /loyalty/vouchers/redeem",
    "Authorization: Bearer <SHOP_ADMIN token>",
    "",
    "{",
    '  "code":     "VCH-AB12-CD34-EF56",',
    '  "shopId":   "5b1c2d3e-4567-890a-bcde-f01234567890",',
    '  "phoneNumber": "+263771234567"     // optional but recommended (anti-fraud)',
    "}",
])
body(doc, "Success response (200 OK):")
code_block(doc, [
    "{",
    '  "code": "200 OK",',
    '  "message": "Voucher redeemed successfully",',
    '  "data": {',
    '    "redemptionId":  "e3d9a1b2-1234-5678-2345-678901234567",',
    '    "voucherId":     "c1b7e9f0-9012-3456-0123-456789012345",',
    '    "status":        "REDEEMED",      // or PARTIALLY_USED for MULTI_USE',
    '    "discountType":  "FIXED",',
    '    "discountValue": 5.00,',
    '    "usesRemaining": 0,',
    '    "redeemedAt":    "2026-06-11T10:28:14"',
    "  }",
    "}",
])
body(doc, "Apply discountValue to the customer's cart before computing the cashAmount you send "
          "to shop-checkout — i.e. the cash leg always reflects what the customer actually paid, "
          "AFTER the voucher discount.")
body(doc, "Failure responses:")
code_block(doc, [
    "# Expired / already redeemed / wrong merchant",
    "{ \"code\": \"400 BAD_REQUEST\", \"message\": \"Voucher is expired\", \"data\": null }",
    "{ \"code\": \"400 BAD_REQUEST\", \"message\": \"Voucher has been fully used\", \"data\": null }",
    "{ \"code\": \"400 BAD_REQUEST\", \"message\": \"Voucher is not valid at this merchant\", \"data\": null }",
    "",
    "# Invalid signature (someone forged the code)",
    "{ \"code\": \"400 BAD_REQUEST\", \"message\": \"Voucher code signature is invalid\", \"data\": null }",
    "",
    "# Velocity threshold tripped (anti-fraud — too many failed attempts)",
    "{ \"code\": \"429 TOO_MANY_REQUESTS\", \"message\": \"Redemption blocked due to fraud velocity\", \"data\": null }",
])

heading(doc, "7.3 QR-code redemption (alternative)", 2)
body(doc, "If the customer presents a QR code rather than a written code, use the QR endpoint "
          "instead — it accepts the scanner output directly:")
code_block(doc, [
    "POST /loyalty/qr/consume",
    "Authorization: Bearer <SHOP_ADMIN token>",
    "",
    "{",
    '  "qrPayload": "eyJ2IjoxLCJjIjoiVkNILUFCMTItQ0QzNC1FRjU2IiwibiI6IjU3MyJ9.signature",',
    '  "shopId":    "5b1c2d3e-4567-890a-bcde-f01234567890"',
    "}",
])
body(doc, "QR codes carry an embedded nonce and have a short TTL (5 minutes by default) — once "
          "consumed, the same QR cannot be reused, even within its TTL window. The response "
          "shape matches /loyalty/vouchers/redeem.")

heading(doc, "7.4 Direct voucher issuance (CRM)", 2)
body(doc, "If you want to gift a specific customer a one-off voucher (apology, retention, VIP), "
          "issue one from an existing template at MERCHANT_ADMIN level:")
code_block(doc, [
    "POST /loyalty/vouchers/issue",
    "Authorization: Bearer <MERCHANT_ADMIN token>",
    "",
    "{",
    '  "templateId":  "11111111-2222-3333-4444-555555555555",',
    '  "phoneNumber": "+263771234567",       // optional — leave null for unassigned',
    '  "note":        "Goodwill replacement for order #4892"',
    "}",
])
body(doc, "Response carries the voucher code and an expiresAt — share the code with the "
          "customer over SMS/WhatsApp; they can then redeem it via §7.2 like any other voucher.")

# ----------------------- 8. Wallet
doc.add_page_break()
heading(doc, "8. Customer wallet lookup", 1)
body(doc, "If your POS surfaces \"current balance\" before applying points, query the wallet directly:")
code_block(doc, [
    "GET /loyalty/users/me/wallet",
    "Authorization: Bearer <CUSTOMER token>",
])
body(doc, "Response:")
code_block(doc, [
    "{",
    '  "code": "200 OK",',
    '  "message": "Wallet retrieved",',
    '  "data": {',
    '    "phoneNumber":   "+263771234567",',
    '    "totalPoints":   1812.5000,',
    '    "totalVouchers": 3',
    "  }",
    "}",
])
body(doc, "This endpoint requires a CUSTOMER token (so the customer's own app must be the "
          "caller; SHOP_ADMIN cannot read other people's wallets via /me). For a till-side "
          "balance check, use the voucher-list endpoint in §7.1 instead — it returns the "
          "customer's active vouchers and is enough information for the cashier to compute "
          "the maximum points the customer could redeem.")

# ----------------------- 9. Error handling
heading(doc, "9. Error handling — what to do per status code", 1)
body(doc, "All errors return the same envelope: { \"code\": \"…\", \"message\": \"…\", \"data\": null }.")
endpoint_summary_table(doc, [
    ("400", "BAD_REQUEST", "Validation, business-rule violation, expired voucher, insufficient balance.", "Surface message to cashier; do not retry."),
    ("401", "UNAUTHORIZED", "Missing, invalid, or expired JWT.", "Refresh token; re-login if refresh fails."),
    ("403", "FORBIDDEN", "Authenticated but wrong role (e.g. SHOP_ADMIN calling /loyalty/merchants).", "Re-check role; not retryable."),
    ("404", "NOT_FOUND", "Shop, merchant, voucher, or customer not in this tenant.", "Verify IDs; not retryable."),
    ("409", "CONFLICT", "Idempotency collision, e.g. the same reference already exists.", "Generate a fresh reference, then retry."),
    ("429", "TOO_MANY_REQUESTS", "Rate limit (50 req/s normal, 100 burst per token).", "Back off ~1 s; honor the Retry-After header."),
    ("502", "BAD_GATEWAY", "Downstream service returned a fault.", "Retry once after ~1 s."),
    ("503", "SERVICE_UNAVAILABLE", "Loyalty service unreachable.", "Retry once after ~500 ms; queue if it persists."),
])

# ----------------------- 10. End-to-end example
doc.add_page_break()
heading(doc, "10. End-to-end example", 1)
body(doc, "Tanaka (phone +263771234567) walks into the Avondale branch of Chicken Inn for the "
          "first time, orders a $15 main meal, and uses a $5 voucher she got via SMS. This is "
          "every call the POS makes, in order.")

heading(doc, "Step 1 — Operator logs in (once per shift)", 3)
code_block(doc, [
    "POST /auth/login",
    '{ "identifier": "till-1@chickeninn.co.zw", "password": "…" }',
    "→ 200 OK, captures token = eyJ…",
])

heading(doc, "Step 2 — Tanaka is new; enrol her", 3)
code_block(doc, [
    "POST /auth/customer/register",
    '{ "firstName": "Tanaka", "lastName": "Moyo",',
    '  "phoneNumber": "+263771234567", "country": "Zimbabwe" }',
    "→ 201 Created",
])

heading(doc, "Step 3 — List her active vouchers", 3)
code_block(doc, [
    "GET /loyalty/vouchers/users/by-phone/+263771234567/active",
    "Authorization: Bearer eyJ…",
    "→ 200 OK, returns [{ code: \"VCH-AB12-CD34-EF56\", value: $5 FIXED }, …]",
])

heading(doc, "Step 4 — Redeem the $5 voucher she chose", 3)
code_block(doc, [
    "POST /loyalty/vouchers/redeem",
    '{ "code": "VCH-AB12-CD34-EF56", "shopId": "…", "phoneNumber": "+263771234567" }',
    "→ 200 OK, status \"REDEEMED\", discountValue 5.00",
])

heading(doc, "Step 5 — Post the sale ($15 − $5 = $10)", 3)
code_block(doc, [
    "POST /payments/shop-checkout",
    "X-Customer-Phone: +263771234567",
    '{ "shopId": "…", "paymentMethod": "CASH", "cashAmount": 10.00 }',
    "→ 200 OK, pointsEarned 12.5000, walletBalanceAfter 12.5000,",
    "          reference \"SHOP-7c9e6679-…\"",
])

heading(doc, "Step 6 — Print the receipt", 3)
body(doc, "Print: subtotal $15.00 · voucher discount −$5.00 · paid $10.00 · points earned 12.50 · "
          "balance 12.50 · ref SHOP-7c9e6679-… Reuse the same reference if you ever need to "
          "void / reverse the sale (a future API).")

# ----------------------- 11. Sandbox + production
heading(doc, "11. Base URLs and rollout", 1)
body(doc, "Two environments. Use sandbox for integration; switch the base URL to production at go-live.")
endpoint_summary_table(doc, [
    ("Sandbox",    "https://sandbox.dtx.innbucks.co.zw", "Test merchants and shops; data wiped weekly.", "Open to any partner"),
    ("Production", "https://dtx.innbucks.co.zw",          "Live customer wallets; signed go-live form required.", "Approved partners only"),
])
body(doc, "Rate limits in production: 50 requests/second sustained, 100 burst, per access "
          "token. Sandbox is more permissive (1000/2000) so you can stress-test integration.")

# ----------------------- 12. Support
heading(doc, "12. Support", 1)
body(doc, "Integration questions, sandbox provisioning, and go-live sign-off:")
sup = doc.add_paragraph()
add_run(sup, "Tawanda Mpofu — ", bold=True, color=BRAND_BLUE, size=10.5)
add_run(sup, "Senior Backend Engineer · Innovation Team", color=TEXT_DARK, size=10.5)
sup.add_run().add_break()
add_run(sup, "tmpofu@innbucks.co.zw", color=BRAND_BLUE, size=10.5)
sup.add_run().add_break()
add_run(sup, "Please include the offending request id (X-Correlation-Id header) and an exact "
             "timestamp on any production support ticket.", color=TEXT_MUTED, italic=True, size=9.5)

# Footer / page numbers
section = doc.sections[0]
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer, "InnBucks · Confidential · Merchant POS Integration Guide v1.0", color=TEXT_MUTED, size=8)

out = Path(__file__).resolve().parent / "InnBucks-Merchant-POS-Integration-Guide.docx"
doc.save(out)
print(f"Wrote {out}")
