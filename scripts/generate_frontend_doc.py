"""Generate the LVMP frontend integration guide as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- styles ----------------------------------------------------------------
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


# --- content ---------------------------------------------------------------
title = doc.add_heading("LVMP Backend → Frontend Integration Guide", 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0B, 0x2A, 0x4A)
add_para(
    "A walkthrough of how the loyalty-service exposes its capabilities and how a "
    "SuperApp / admin portal should consume them.",
    italic=True,
)
doc.add_paragraph()

# 1
add_heading("1. Base URL & multi-tenancy", 1)
add_para("All loyalty endpoints live under the API gateway:")
add_code("https://<gateway-host>/api/loyalty/...")
add_para(
    "Every request must carry a tenant header (the gateway forwards it to the service):",
    bold=True,
)
add_code("X-Tenant-Id:   <tenant-uuid>\n# or\nX-Tenant-Code: ACME")
add_para(
    "A request without one returns 400 MISSING_TENANT. A request with a tenant "
    "the caller doesn't belong to returns 403 CROSS_TENANT."
)
add_para("The gateway authenticates the caller via the existing JWT flow on /auth/**.")
add_code("Authorization: Bearer <jwt>")

# 2
add_heading("2. Standard response shapes", 1)
add_para("Success — the resource as JSON.")
add_para("Error — a uniform envelope:")
add_code(
    '{\n'
    '  "timestamp": "2026-05-03T14:00:00Z",\n'
    '  "status": 400,\n'
    '  "code": "INSUFFICIENT_FUNDS",\n'
    '  "message": "wallet balance would go negative"\n'
    '}'
)
add_para("Validation failures also include a fields map:")
add_code('{ "code": "VALIDATION", "fields": { "phone": "must not be blank" } }')
add_para(
    "code is stable and machine-friendly; show message to operators, branch on "
    "code in the UI."
)

# 3
add_heading("3. Lifecycle: who creates what, in what order", 1)

add_heading("a) Operator (platform admin) onboards a tenant", 2)
add_code('POST /api/loyalty/tenants\n{ "code": "ACME", "name": "Acme Holdings" }')
add_para("Returns the new tenant.id. Subsequent requests carry X-Tenant-Code: ACME.")

add_heading("b) Tenant admin onboards merchants", 2)
add_code(
    'POST /api/loyalty/merchants\n'
    '{\n'
    '  "name": "Cafe Downtown",\n'
    '  "category": "F&B",\n'
    '  "location": "Harare CBD",\n'
    '  "currency": "USD",\n'
    '  "billingCycle": "MONTHLY",\n'
    '  "feePerPointIssued": "0.001",\n'
    '  "feePerVoucherIssued": "0.05",\n'
    '  "feePerVoucherRedeemed": "0.10"\n'
    '}'
)
add_para("Fee fields drive periodic invoicing. They can be zero.")

add_heading("c) Tenant admin defines loyalty rules", 2)
add_code(
    'POST /api/loyalty/rules\n'
    '{\n'
    '  "merchantId": "<merchant-uuid or null for global>",\n'
    '  "transactionType": "PURCHASE",\n'
    '  "pointsPerUnit": "1",\n'
    '  "multiplier": "1",\n'
    '  "maxPointsPerTxn": null,\n'
    '  "pocket": null\n'
    '}'
)
add_para(
    "merchantId: null = global template. A merchant-specific row beats the global one. "
    "Time-bound campaigns layer on top:"
)
add_code(
    'POST /api/loyalty/rules/campaigns\n'
    '{\n'
    '  "merchantId": "<uuid>",\n'
    '  "name": "Weekend 2x",\n'
    '  "multiplier": "2",\n'
    '  "transactionType": "PURCHASE",\n'
    '  "startsAt": "2026-05-10T00:00:00Z",\n'
    '  "endsAt":   "2026-05-12T23:59:59Z"\n'
    '}'
)

add_heading("d) End users register", 2)
add_code(
    'POST /api/loyalty/users\n'
    '{ "phone": "+263770000001", "fullName": "Alice", "country": "ZW", '
    '"merchantId": "<optional>" }'
)
add_para(
    "A MAIN wallet is created automatically. The frontend should cache the "
    "returned user.id."
)

# 4
add_heading("4. The two flows the SuperApp needs most", 1)

add_heading("Flow A — Earn points from a transaction", 2)
add_code(
    'POST /api/loyalty/transactions\n'
    '{\n'
    '  "merchantId": "<uuid>",\n'
    '  "userId": "<uuid>",\n'
    '  "type": "PURCHASE",\n'
    '  "amount": "100.00",\n'
    '  "currency": "USD",\n'
    '  "reference": "POS-INV-12345"\n'
    '}'
)
add_para("Response:")
add_code(
    '{\n'
    '  "id": "<txn-uuid>",\n'
    '  "type": "PURCHASE",\n'
    '  "amount": "100.00",\n'
    '  "pointsDelta": "100.0000",\n'
    '  "balanceAfter": "100.0000",\n'
    '  "ruleId": "<uuid>",\n'
    '  "campaignId": null,\n'
    '  "reference": "POS-INV-12345",\n'
    '  "createdAt": "..."\n'
    '}'
)
add_para("Important rules the frontend must respect:", bold=True)
add_bullet(
    "reference is the merchant's invoice / receipt id. Replaying the same "
    "(merchantId, reference) returns 409 DUPLICATE_REFERENCE — that's the "
    "idempotency contract. Always generate one client-side."
)
add_bullet(
    "pointsDelta = 0 is a valid response (no rule matched / type excluded). "
    "Show \"earned 0 points\" rather than treating it as an error."
)
add_bullet("balanceAfter is the wallet balance for display.")

add_heading("Flow B — Redeem a voucher", 2)
add_code(
    'POST /api/loyalty/vouchers/redeem\n'
    '{\n'
    '  "code": "Q7K2NB3FXP4M",\n'
    '  "userId": "<optional, if logged in>",\n'
    '  "merchantId": "<uuid>",\n'
    '  "outletCode": "BR-001",\n'
    '  "deviceFingerprint": "<stable per device>",\n'
    '  "ipAddress": "<auto, optional>"\n'
    '}'
)
add_para("Possible outcomes the UI must handle:")
add_table(
    [
        ["200", "—", "Success: show value redeemed (value, valueType, status, usesRemaining)"],
        ["404", "NOT_FOUND", "Invalid voucher code"],
        ["403", "BAD_SIGNATURE", "Invalid voucher code (don't leak that signature failed)"],
        ["400", "EXPIRED", "This voucher has expired"],
        ["409", "ALREADY_REDEEMED", "This voucher has already been used"],
        ["409", "REVOKED", "This voucher is no longer valid"],
        ["403", "WRONG_MERCHANT", "Not redeemable here"],
        ["403", "USER_BLOCKED", "Account locked — contact support"],
    ],
    ["HTTP", "code", "What to show"],
)
add_para(
    "Every rejection is logged in fraud_attempts server-side; after N failures from "
    "the same deviceFingerprint within a window, the user is auto-blocked. This is "
    "invisible to the UI except that further attempts return USER_BLOCKED."
)

# 5
add_heading("5. Voucher lifecycle the SuperApp visualises", 1)
add_code(
    "ISSUED → DELIVERED → VIEWED → REDEEMED\n"
    "                              ↘ PARTIALLY_USED  (multi-use)\n"
    "                              ↘ EXPIRED         (after expiresAt)\n"
    "                              ↘ REVOKED         (admin action)"
)
add_para("When the user opens a voucher card, fire-and-forget:")
add_code("POST /api/loyalty/vouchers/codes/{code}/viewed")
add_para(
    "That moves the voucher from DELIVERED to VIEWED, which the operator "
    "dashboard uses for engagement metrics."
)
add_para("To list a user's active vouchers:")
add_code("GET /api/loyalty/vouchers/users/{userId}/active")
add_para(
    "Returns vouchers in any of ISSUED | DELIVERED | VIEWED | PARTIALLY_USED. "
    "Use this for the \"My Vouchers\" tab."
)

# 6
add_heading("6. Wallets, sub-wallets, transfers", 1)
add_para(
    "Every user has one main wallet automatically. They can create additional "
    "sub-wallets (\"pockets\") for specific reward types or savings goals:"
)
add_code(
    'POST /api/loyalty/users/{id}/wallets\n'
    '{ "label": "Family pocket", "pocket": "FAMILY", "type": "SUB", "lockedUntil": null }'
)
add_para(
    "type=SAVINGS + lockedUntil=2027-01-01 produces a wallet that can't be debited "
    "until that date — useful for goal-saving UX."
)
add_para("Peer-to-peer transfer (subject to sender having balance):")
add_code(
    'POST /api/loyalty/transfer\n'
    '{\n'
    '  "fromUserId": "<uuid>",\n'
    '  "toUserId":   "<uuid>",\n'
    '  "points":     "20",\n'
    '  "reason":     "Birthday gift"\n'
    '}'
)
add_para(
    'Response: { "status": "OK", "newSenderBalance": "30.0000" }. '
    "Failure modes: BAD_AMOUNT, SELF_TRANSFER, INSUFFICIENT_FUNDS."
)

# 7
add_heading("7. The unified SuperApp dashboard", 1)
add_para("One call gives you everything for the home screen:")
add_code("GET /api/loyalty/reports/user/{userId}")
add_para("Returns:")
add_code(
    '{\n'
    '  "userId": "...",\n'
    '  "totalPoints": "450.00",\n'
    '  "wallets": [ { "id": "...", "label": "Main", "type": "MAIN", "balance": "300.00" }, ... ],\n'
    '  "activeVouchers": [ { "code": "Q7K2N...", "status": "DELIVERED", "expiresAt": "..." }, ... ],\n'
    '  "recentTransactions": [ { "type": "PURCHASE", "pointsDelta": "100", "createdAt": "..." }, ... ]\n'
    '}'
)
add_para("Render four cards from this single payload — no waterfall calls.")

# 8
add_heading("8. QR flows", 1)

add_heading("Merchant-issued QR (customer earns)", 2)
add_para("The merchant terminal calls:")
add_code(
    'POST /api/loyalty/qr/issue\n'
    '{\n'
    '  "sourceType": "MERCHANT",\n'
    '  "sourceId":   "<merchant-uuid>",\n'
    '  "transactionType": "PURCHASE",\n'
    '  "amount":   "50",\n'
    '  "currency": "USD",\n'
    '  "ttlSeconds": 300\n'
    '}'
)
add_para("Response:")
add_code(
    '{\n'
    '  "token":     "abc123...",\n'
    '  "signature": "f3a9...",\n'
    '  "tenantId":  "...",\n'
    '  "sourceType":"MERCHANT",\n'
    '  "sourceId":  "...",\n'
    '  "transactionType":"PURCHASE",\n'
    '  "expiresAt": "..."\n'
    '}'
)
add_para("Encode the JSON in the QR. The SuperApp scans and posts:")
add_code(
    'POST /api/loyalty/qr/consume\n'
    '{\n'
    '  "token":     "abc123...",\n'
    '  "signature": "f3a9...",\n'
    '  "userId":    "<scanning-user-uuid>",\n'
    '  "reference": "<idempotency key>"\n'
    '}'
)
add_para(
    "Result: a TransactionResponse (same shape as Flow A), and points land in "
    "the user's main wallet."
)

add_heading("User-issued QR (peer-to-peer transfer)", 2)
add_para(
    "Same qr/issue endpoint with sourceType=USER, sourceId=<senderUserId>, "
    "transactionType=TRANSFER, amount=<points>. The receiving user scans and "
    "posts to qr/consume — internally this routes to the transfer service."
)
add_para(
    "Errors the scanner UI must handle:", bold=True
)
add_para(
    "QR_EXPIRED, QR_REUSED, BAD_SIGNATURE. All three should show the same generic "
    "\"QR no longer valid\" message; logs already capture the specifics."
)

# 9
add_heading("9. Admin dashboards", 1)
add_para("Three role-scoped read-only endpoints:")
add_code(
    "GET /api/loyalty/reports/operator         # platform-wide\n"
    "GET /api/loyalty/reports/tenant           # tenant-scoped (header-bound)\n"
    "GET /api/loyalty/reports/merchant/{id}    # merchant-scoped"
)
add_para("Plus filter / export utilities:")
add_code(
    "GET /api/loyalty/reports/transactions/mix?merchantId=<>&from=2026-05-01&to=2026-05-31\n"
    "GET /api/loyalty/reports/transactions/export?merchantId=<>&from=...&to=...   # text/csv\n"
    "GET /api/loyalty/reports/fraud"
)
add_para(
    "Each returns flat JSON ready for charts; the operator dashboard payload "
    "includes \"expiring in 7/30 days\" counts so the UI can highlight expiring "
    "inventory without a second call."
)

# 10
add_heading("10. Invoicing (merchant finance UI)", 1)
add_code(
    "GET  /api/loyalty/invoices/merchant/{merchantId}    # list\n"
    "POST /api/loyalty/invoices/{id}/pay                 # mark paid\n"
    "POST /api/loyalty/invoices/generate                 # ad-hoc"
)
add_para(
    "Body for generate: { merchantId, periodStart, periodEnd }. The scheduler "
    "auto-runs once a day and emits one PENDING invoice per active merchant for "
    "the previous billing period. The UI just lists and toggles status; "
    "aggregation already happened on the server."
)

# 11
add_heading("11. Mini-apps manifest", 1)
add_para("Cross-platform shells fetch their available mini-apps with:")
add_code(
    "GET /api/loyalty/mini-apps/manifest?merchantId=<optional>\n"
    "→ [ { \"slug\":\"fuel-rewards\", \"name\":\"...\", \"iconUrl\":\"...\", \"entryUrl\":\"...\" }, ... ]"
)
add_para(
    "Render each as a tile on the SuperApp grid; the tile opens entryUrl "
    "(a webview or deep link)."
)

# 12
add_heading("12. Things to NOT call", 1)
add_code(
    "POST /api/loyalty/convert-to-airtime   →  always returns\n"
    '{ "status":"NOT_ENABLED",\n'
    '  "message":"M-Pesa / airtime conversion is not enabled in this build." }'
)
add_para(
    "Hide the button until the backend signals integration is live (we'll add a "
    "feature flag later)."
)

# 13
add_heading("13. Practical sequencing for the first SuperApp screen", 1)
add_code(
    "on app launch:\n"
    "  GET /auth/me                              # existing user-service\n"
    "  GET /api/loyalty/reports/user/{me.id}      # paint wallets + vouchers + history\n"
    "\n"
    "on \"Scan to earn\":\n"
    "  scan QR → POST /api/loyalty/qr/consume     # refresh dashboard\n"
    "\n"
    "on \"My vouchers\" tap:\n"
    "  list comes from cached dashboard payload\n"
    "\n"
    "on \"Redeem\" inside a voucher:\n"
    "  POST /api/loyalty/vouchers/codes/{code}/viewed   # eager, fire-and-forget\n"
    "  if user is at a merchant POS → it scans / types and submits redeem itself"
)
add_para(
    "That's the full surface. The backend OpenAPI doc renders at "
    "/loyalty-service/v3/api-docs (and inside the gateway's aggregated Swagger UI), "
    "so the frontend team can generate types + clients directly from there."
)

doc.save("/home/user/ticketing-system/docs/LVMP_Frontend_Integration_Guide.docx")
print("Wrote /home/user/ticketing-system/docs/LVMP_Frontend_Integration_Guide.docx")
